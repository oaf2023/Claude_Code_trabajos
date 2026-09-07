# ============================================================================
# Archivo         : verificar_documentos.py
# Descripcion     : Verifica los .docx generados: apertura, conteo de
#                   parrafos/tablas, busqueda de posibles secretos y de
#                   marcadores pendientes, y estadisticas de cobertura.
# Autor           : Equipo SafeAlert
# Fecha           : 2026-09-06
# Version         : 1.0.0
# Lenguaje        : Python 3.13
# Uso             : python verificar_documentos.py <carpeta_documentos>
# ============================================================================

import glob
import os
import re
import sys

from docx import Document

PATRONES_SECRETOS = [
    (r"AIza[0-9A-Za-z\-_]{35}", "Clave API Google/Firebase"),
    (r"AKIA[0-9A-Z]{16}", "Clave AWS"),
    (r"sk_live_[0-9A-Za-z]{10,}", "Clave Stripe live"),
    (r"sk_test_[0-9A-Za-z]{10,}", "Clave Stripe test"),
    (r"ghp_[0-9A-Za-z]{30,}", "Token GitHub"),
    (r"xox[baprs]-[0-9A-Za-z\-]{10,}", "Token Slack"),
    (r"-----BEGIN [A-Z ]*PRIVATE KEY-----", "Clave privada"),
    (r"AC[0-9a-f]{32}", "Twilio SID"),
    (r"SG\.[0-9A-Za-z\-_]{20,}", "Clave SendGrid"),
    (r"eyJhbGciOi[A-Za-z0-9+/=]{10,}", "JWT"),
    (r"(?:secret|password|token|api[_-]?key)\s*[:=]\s*['\"]?[A-Za-z0-9_\-\.]{12,}",
     "Posible asignacion de secreto"),
]


def texto_completo(doc):
    partes = []
    for p in doc.paragraphs:
        partes.append(p.text)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.append(celda.text)
    return "\n".join(partes)


def main():
    carpeta = sys.argv[1] if len(sys.argv) > 1 else r"C:\Claude_Code_trabajos\safealert\documentacion_generada"
    docx = sorted(glob.glob(os.path.join(carpeta, "*.docx")) +
                  glob.glob(os.path.join(carpeta, "anexos", "*.docx")))
    print("Documentos a verificar: %d\n" % len(docx))
    total_parrafos = 0
    total_tablas = 0
    total_palabras = 0
    problemas = 0
    for ruta in docx:
        try:
            doc = Document(ruta)
        except Exception as exc:  # noqa: BLE001
            print("[FALLO APERTURA] %s -> %r" % (ruta, exc))
            problemas += 1
            continue
        np = len(doc.paragraphs)
        nt = len(doc.tables)
        txt = texto_completo(doc)
        palabras = len(txt.split())
        total_parrafos += np
        total_tablas += nt
        total_palabras += palabras
        # Secretos
        hallazgos = []
        for patron, nombre in PATRONES_SECRETOS:
            for m in re.finditer(patron, txt, re.IGNORECASE):
                hallazgos.append((nombre, m.group(0)[:40]))
        estado = "OK"
        if hallazgos:
            estado = "REVISAR"
            problemas += 1
            for nombre, muestra in hallazgos[:6]:
                print("   [POSIBLE SECRETO] %s -> %s..." % (nombre, muestra))
        # Marcadores pendientes sin resolver
        pend = re.findall(r"\[(?:TODO|FIXME|XXX|PENDIENTE)\]", txt)
        print("%-70s parrafos=%6d tablas=%3d palabras=%8d %s"
              % (os.path.basename(ruta), np, nt, palabras, estado))
        if hallazgos:
            print("   Total posibles secretos: %d" % len(hallazgos))
        if pend:
            print("   Marcadores TODO/FIXME en texto: %d" % len(pend))
    print("\nTOTAL parrafos=%d tablas=%d palabras=%d problemas=%d"
          % (total_parrafos, total_tablas, total_palabras, problemas))
    return 0


if __name__ == "__main__":
    sys.exit(main())
