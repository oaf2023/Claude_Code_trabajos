# ============================================================================
# Archivo         : md_docx.py
# Descripcion     : Convierte el subconjunto Markdown definido en
#                   CONVENCIONES_MD.md a documentos Microsoft Word (.docx)
#                   usando python-docx. Estilos corporativos, cabecera y
#                   pie de pagina "Pagina X de Y".
# Autor           : Equipo SafeAlert
# Fecha           : 2026-09-06
# Version         : 1.0.0
# Lenguaje        : Python 3.13 (python-docx 1.2.0)
# Uso             : import md_docx; md_docx.anadir_markdown(documento, md)
# ============================================================================

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Marcadores de parrafo admitidos y su color (RGB)
MARCADORES = [
    ("[CRITICO]", RGBColor(0xB0, 0x00, 0x00)),
    ("[CRÍTICO]", RGBColor(0xB0, 0x00, 0x00)),
    ("[ALTO]", RGBColor(0xC0, 0x00, 0x00)),
    ("[RIESGO]", RGBColor(0xC0, 0x39, 0x2B)),
    ("[ADVERTENCIA]", RGBColor(0xBF, 0x6A, 0x00)),
    ("[MEDIO]", RGBColor(0xB8, 0x6E, 0x00)),
    ("[BAJO]", RGBColor(0x1F, 0x6F, 0xB0)),
    ("[INFORMATIVO]", RGBColor(0x1F, 0x6F, 0xB0)),
    ("[NOTA]", RGBColor(0x1F, 0x6F, 0xB0)),
    ("[OBSERVACION TECNICA]", RGBColor(0x6A, 0x2D, 0x9E)),
    ("[OBSERVACIÓN TÉCNICA]", RGBColor(0x6A, 0x2D, 0x9E)),
    ("[POTENCIALMENTE NO UTILIZADO]", RGBColor(0x6A, 0x2D, 0x9E)),
    ("[RECOMENDACION]", RGBColor(0x1E, 0x7A, 0x46)),
    ("[RECOMENDACIÓN]", RGBColor(0x1E, 0x7A, 0x46)),
    ("[NIVEL DE CERTEZA", RGBColor(0x45, 0x55, 0x64)),
    ("[SECRETO OCULTO]", RGBColor(0x6A, 0x2D, 0x9E)),
]

FUENTE_CODIGO = "Consolas"
FUENTE_TEXTO = "Calibri"
TAM_CODIGO = Pt(8)
TAM_TEXTO = Pt(10.5)


def _configurar_estilos(doc):
    """Configura estilos base del documento."""
    normal = doc.styles["Normal"]
    normal.font.name = FUENTE_TEXTO
    normal.font.size = TAM_TEXTO
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    # Titulos con color corporativo rojo
    for nombre, tam in (("Heading 1", 20), ("Heading 2", 16),
                        ("Heading 3", 13), ("Heading 4", 11.5)):
        estilo = doc.styles[nombre]
        estilo.font.name = FUENTE_TEXTO
        estilo.font.size = Pt(tam)
        estilo.font.bold = True
        estilo.font.color.rgb = RGBColor(0x9A, 0x1B, 0x1B)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)


def _sombreado_parrafo(parrafo, hexcolor):
    """Aplica color de fondo a un parrafo (w:shd)."""
    pPr = parrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def _bordes_celda(celda):
    """Bordes finos grises para celdas de tabla."""
    tcPr = celda._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + lado)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "B0B0B0")
        borders.append(el)
    tcPr.append(borders)


def _sombreado_celda(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _marca_color(texto):
    """Devuelve (color, resto) si el parrafo comienza con un marcador."""
    for token, color in MARCADORES:
        if texto.startswith(token):
            if token.startswith("[NIVEL DE CERTEZA"):
                idx = texto.find("]")
                return color, texto
            return color, texto
    return None, texto


def _anadir_run_rich(parrafo, texto, base_negrita=False, tam=None,
                      fuente=None, color=None):
    """Añade texto con formato inline (**bold**, *italic*, `code`)."""
    # Tokenizacion simple: reemplaza marcadores por placeholders
    partes = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]*`)", texto)
    for parte in partes:
        if not parte:
            continue
        run = parrafo.add_run()
        if parte.startswith("**") and parte.endswith("**") and len(parte) > 3:
            run.text = parte[2:-2]
            run.bold = True
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            run.text = parte[1:-1]
            run.italic = True
        elif parte.startswith("`") and parte.endswith("`") and len(parte) > 1:
            run.text = parte[1:-1]
            run.font.name = FUENTE_CODIGO
            run.font.size = TAM_CODIGO
        else:
            run.text = parte
        if base_negrita:
            run.bold = True
        if tam:
            run.font.size = tam
        if fuente and run.font.name != FUENTE_CODIGO:
            run.font.name = fuente
        if color:
            run.font.color.rgb = color


def _es_fila_separador(linea):
    """True si la linea es separador de tabla tipo | --- | --- |."""
    cuerpo = linea.strip().strip("|")
    return bool(cuerpo) and all(re.fullmatch(r"\s*:?-{3,}:?\s*", c) for c in cuerpo.split("|"))


def _parsear_tabla(lineas, i):
    """Consume lineas de tabla desde i; devuelve (filas, nueva_i)."""
    filas = []
    while i < len(lineas):
        linea = lineas[i]
        if not linea.strip().startswith("|"):
            break
        if _es_fila_separador(linea):
            i += 1
            continue
        celdas = [c.strip() for c in linea.strip().strip("|").split("|")]
        filas.append(celdas)
        i += 1
    return filas, i


def _anadir_tabla(doc, filas):
    """Crea tabla con estilo 'Table Grid' y encabezado sombreado."""
    ncols = max(len(f) for f in filas) if filas else 1
    ncols = min(ncols, 12)
    tabla = doc.add_table(rows=len(filas), cols=ncols)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, fila in enumerate(filas):
        for c in range(ncols):
            celda = tabla.cell(r, c)
            texto = fila[c] if c < len(fila) else ""
            celda.paragraphs[0].text = ""
            p = celda.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            _anadir_run_rich(p, texto)
            if r == 0:
                _sombreado_celda(celda, "F2DCDA")
                for run in p.runs:
                    run.bold = True
            for run in p.runs:
                run.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _anadir_codigo(doc, codigo, lenguaje=""):
    """Añade bloque de codigo en Consolas con fondo claro."""
    parrafo = doc.add_paragraph()
    parrafo.paragraph_format.space_before = Pt(4)
    parrafo.paragraph_format.space_after = Pt(6)
    parrafo.paragraph_format.line_spacing = 1.0
    _sombreado_parrafo(parrafo, "F5F5F5")
    for n, linea in enumerate(codigo):
        if n > 0:
            parrafo.add_run().add_break()
        run = parrafo.add_run(linea if linea else " ")
        run.font.name = FUENTE_CODIGO
        run.font.size = TAM_CODIGO
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FUENTE_CODIGO)
        rFonts.set(qn("w:hAnsi"), FUENTE_CODIGO)
    # Ajustar tamano de parrafo/caracter asiatico
    if not codigo:
        parrafo.add_run(" ")


def _anadir_parrafo(doc, texto, es_cita=False, es_lista=False,
                    es_numerada=False, nivel=0):
    p = doc.add_paragraph()
    if es_lista:
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.left_indent = Cm(0.75 + 0.5 * nivel)
    elif es_numerada:
        p.style = doc.styles["List Number"]
        p.paragraph_format.left_indent = Cm(0.75 + 0.5 * nivel)
    else:
        p.paragraph_format.space_after = Pt(6)
    color, resto = _marca_color(texto)
    if es_cita:
        color = RGBColor(0x60, 0x60, 0x60)
    if color:
        _anadir_run_rich(p, resto, base_negrita=True, color=color)
    elif es_cita:
        _anadir_run_rich(p, texto, color=color)
        for run in p.runs:
            run.italic = True
    else:
        _anadir_run_rich(p, texto)
    return p


def anadir_markdown(doc, texto_md):
    """Convierte un texto Markdown (subconjunto permitido) y lo anade a doc."""
    _configurar_estilos(doc)
    lineas = texto_md.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    n = len(lineas)
    lista_pendiente = None  # 'bullets' | 'numbers'
    while i < n:
        linea = lineas[i]
        t = linea.rstrip()

        # Regla horizontal
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", t.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "9A1B1B")
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        # Encabezado
        m = re.match(r"^(#{1,6})\s+(.*)$", t)
        if m:
            nivel = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=nivel)
            i += 1
            continue

        # Cita
        if t.startswith(">"):
            _anadir_parrafo(doc, t.lstrip(">").strip(), es_cita=True)
            i += 1
            continue

        # Inicio de cerca de codigo
        if t.startswith("```"):
            lenguaje = t[3:].strip()
            bloque = []
            i += 1
            while i < n and not lineas[i].strip().startswith("```"):
                bloque.append(lineas[i].rstrip())
                i += 1
            if i < n:
                i += 1  # saltar cierre
            _anadir_codigo(doc, bloque, lenguaje)
            continue

        # Tabla
        if t.strip().startswith("|"):
            filas, i = _parsear_tabla(lineas, i)
            if filas:
                _anadir_tabla(doc, filas)
            continue

        # Lista de items (bullets)
        mb = re.match(r"^(\s*)[-*]\s+(.*)$", t)
        mn = re.match(r"^(\s*)\d+[.)]\s+(.*)$", t)
        if mb or mn:
            sangria = len((mb or mn).group(1))
            nivel = sangria // 2
            texto_item = (mb or mn).group(2)
            if mb:
                _anadir_parrafo(doc, texto_item, es_lista=True, nivel=nivel)
            else:
                _anadir_parrafo(doc, texto_item, es_numerada=True, nivel=nivel)
            i += 1
            continue

        # Linea vacia -> separacion
        if not t.strip():
            i += 1
            continue

        # Parrafo normal (posible multilinea hasta linea vacia)
        parrafo_lineas = [t]
        i += 1
        while i < n:
            sig = lineas[i].rstrip()
            if (not sig.strip() or sig.startswith("#") or
                    sig.startswith("```") or sig.strip().startswith("|") or
                    re.match(r"^(\s*)[-*]\s+", sig) or
                    re.match(r"^(\s*)\d+[.)]\s+", sig) or
                    re.fullmatch(r"-{3,}", sig.strip()) or sig.startswith(">")):
                break
            parrafo_lineas.append(sig)
            i += 1
        _anadir_parrafo(doc, " ".join(parrafo_lineas))
    return doc


def cabecera_pie(doc, titulo_corto):
    """Cabecera con nombre de proyecto y pie 'Pagina X de Y'."""
    for sec in doc.sections:
        # Cabecera
        hp = sec.header.paragraphs[0]
        hp.text = ""
        run = hp.add_run(titulo_corto + " — Documentación Técnica")
        run.font.name = FUENTE_TEXTO
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Pie: Pagina X de Y
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        r1 = fp.add_run("Página ")
        r1.font.size = Pt(8)
        fld1 = OxmlElement("w:fldSimple")
        fld1.set(qn("w:instr"), "PAGE")
        fp._p.append(fld1)
        r2 = fp.add_run(" de ")
        r2.font.size = Pt(8)
        fld2 = OxmlElement("w:fldSimple")
        fld2.set(qn("w:instr"), "NUMPAGES")
        fp._p.append(fld2)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
