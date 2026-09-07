# ============================================================================
# Archivo         : construir_documentacion.py
# Descripcion     : Genera los documentos .docx de la auditoria de SafeAlert:
#                   documento maestro + anexos por modulo, a partir de los
#                   .md intermedios y de los JSON de matrices.
# Autor           : Equipo SafeAlert
# Fecha           : 2026-09-06
# Version         : 1.0.0
# Lenguaje        : Python 3.13 (python-docx 1.2.0)
# Uso             : python construir_documentacion.py [--solo anexos|maestro]
# Resultado       : documentacion_generada/DOCUMENTACION_TECNICA_COMPLETA_SAFEALERT.docx
#                   documentacion_generada/anexos/ANEXO_*.docx
# ============================================================================

import glob
import json
import os
import re
import sys

sys.path.insert(0, "C:\\Claude_Code_trabajos\\_toolcheck")
import md_docx  # noqa: E402

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROY = r"C:\Claude_Code_trabajos\safealert"
OUT = os.path.join(PROY, "documentacion_generada")
INT = os.path.join(OUT, "intermediate")
MAT = os.path.join(INT, "matrices")
ANEXOS_DIR = os.path.join(OUT, "anexos")
MASTER_DIR = os.path.join(INT, "master")

ANEXOS = [
    {
        "codigo": "Anexo A",
        "titulo": "ANEXO A — APP MOVIL PRINCIPAL SAFEALERT (EXPO / REACT NATIVE)",
        "archivo": "ANEXO_A_APP_MOVIL_SAFEALERT.docx",
        "carpetas": ["app_nucleo_config", "app_base_tipos",
                     "app_servicios_alerta", "app_servicios_audio",
                     "app_servicios_localizacion", "app_servicios_pago_cuenta",
                     "app_componentes", "app_pantallas_core",
                     "app_pantallas_onboarding"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Este anexo contiene el analisis archivo por archivo, funcion por "
            "funcion y linea por linea del codigo fuente de la aplicacion movil "
            "principal SafeAlert (Expo SDK 55 / React Native 0.83 / TypeScript): "
            "configuracion, tipos, servicios, componentes, pantallas "
            "(expo-router), stores (Zustand), hooks, utilidades y shims web. "
            "Cada capitulo reproduce los bloques de codigo original y explica "
            "las lineas relevantes inmediatamente debajo. Los archivos aparecen "
            "en orden alfabetico de ruta. El indice agregado de funciones, "
            "clases y endpoints de toda la app esta en el documento maestro."
        ),
    },
    {
        "codigo": "Anexo B",
        "titulo": "ANEXO B — APLICACION IPHONE (VARIANTE EXPO-ROUTER)",
        "archivo": "ANEXO_B_APP_IPHONE.docx",
        "carpetas": ["iphone_app"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Analisis completo de la variante de aplicacion contenida en la "
            "carpeta iphone/ (expo-router). Incluye la comparacion con la app "
            "principal y la determinacion del rol real de esta variante en el "
            "repositorio."
        ),
    },
    {
        "codigo": "Anexo C",
        "titulo": "ANEXO C — BACKEND FLASK Y DESPLIEGUE EN CLOUD RUN",
        "archivo": "ANEXO_C_BACKEND_FLASK_CLOUD_RUN.docx",
        "carpetas": ["backend_flask", "backend_resto"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Analisis de la API REST Flask (backend/flask_app.py, wsgi.py), "
            "dependencias, migraciones SQL (esquema MySQL), pruebas y la "
            "configuracion de despliegue en Google Cloud Run (Dockerfile y "
            "cloudbuild.yaml). Se incluye el diccionario de datos de las "
            "tablas."
        ),
    },
    {
        "codigo": "Anexo D",
        "titulo": "ANEXO D — CLOUD FUNCTIONS DE FIREBASE (TYPESCRIPT)",
        "archivo": "ANEXO_D_CLOUD_FUNCTIONS_FIREBASE.docx",
        "carpetas": ["funciones_firebase"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Analisis de las Cloud Functions de Firebase en TypeScript: "
            "envio de SMS, ordenes y webhooks de pago (Mercado Pago), "
            "limpieza de alertas antiguas y sincronizacion de usuarios."
        ),
    },
    {
        "codigo": "Anexo E",
        "titulo": "ANEXO E — PANEL DE ADMINISTRACION WEB (REACT + VITE)",
        "archivo": "ANEXO_E_PANEL_ADMIN_WEB.docx",
        "carpetas": ["admin_panel"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Analisis del panel de administracion web (React + Vite + "
            "TypeScript): paginas, componentes, cliente de API, estilos y "
            "configuracion del proyecto."
        ),
    },
    {
        "codigo": "Anexo F",
        "titulo": "ANEXO F — CONFIGURACION RAIZ, SCRIPTS, WEB PWA Y PUBLICACION",
        "archivo": "ANEXO_F_CONFIG_SCRIPTS_WEB_PUBLICACION.docx",
        "carpetas": ["config_raiz", "web_scripts_legacy",
                     "publicar_distribucion"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Analisis del punto de entrada y configuracion raiz del proyecto "
            "(Expo, Metro, Babel, Jest, EAS, Firebase, plugins de configuracion "
            "nativa), la salida web PWA y los scripts raiz, los archivos "
            "legados de diagnostico (diag*.mjs, snapshots de UI, informe "
            "tecnico) y el proceso de publicacion a Google Play (PowerShell y "
            "documentacion de Play Console)."
        ),
    },
    {
        "codigo": "Anexo G",
        "titulo": "ANEXO G — DOCUMENTACION EXISTENTE DEL PROYECTO",
        "archivo": "ANEXO_G_DOCUMENTACION_EXISTENTE.docx",
        "carpetas": ["docs_existentes"],
        "intro": (
            "## Alcance de este anexo\n\n"
            "Resumen y evaluacion de la documentacion Markdown existente en el "
            "repositorio (arquitectura, setup, deploy, tutoriales, politicas, "
            "runbooks), con la coherencia verificada contra el codigo real."
        ),
    },
]

TOKENS_TABLA = {
    "[[FUNCIONES_TABLA]]",
    "[[CLASES_TABLA]]",
    "[[TIPOS_TABLA]]",
    "[[ENDPOINTS_TABLA]]",
    "[[DEPENDENCIAS_TABLA]]",
    "[[ENV_TABLA]]",
    "[[BD_TABLA]]",
    "[[RESUMEN_DIRECTORIOS_TABLA]]",
}


def tabla_markdown(encabezados, filas, max_celda=120):
    """Construye una tabla markdown (subconjunto permitido)."""
    def limpiar(c):
        c = str(c).replace("|", "/").replace("\n", " ").strip()
        if len(c) > max_celda:
            c = c[:max_celda - 1] + "…"
        return c
    lineas = ["| " + " | ".join(limpiar(h) for h in encabezados) + " |"]
    lineas.append("| " + " | ".join("---" for _ in encabezados) + " |")
    for fila in filas:
        celdas = [limpiar(c) for c in fila]
        while len(celdas) < len(encabezados):
            celdas.append("")
        lineas.append("| " + " | ".join(celdas[:len(encabezados)]) + " |")
    return "\n".join(lineas)


def leer_json(nombre):
    ruta = os.path.join(MAT, nombre)
    if not os.path.exists(ruta):
        return None
    with open(ruta, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Utilidades docx nativas (portada, control, indice, pagina)
# ---------------------------------------------------------------------------

def _config_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def anadir_portada_maestro(doc, titulo, subtitulo):
    _config_doc(doc)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SafeAlert")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x9A, 0x1B, 0x1B)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aplicacion de alertas SOS por voz y ubicacion")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size = Pt(22)
    r.font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.font.size = Pt(13)
    for _ in range(4):
        doc.add_paragraph()
    for linea in ("Version: 1.0",
                  "Fecha: 06 de septiembre de 2026",
                  "Fecha de analisis: 06 de septiembre de 2026",
                  "Generado a partir del codigo fuente del repositorio",
                  "Autor: Equipo SafeAlert"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linea)
        r.font.size = Pt(11)
    doc.add_page_break()


def anadir_portada_anexo(doc, titulo, descripcion):
    _config_doc(doc)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SafeAlert")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x9A, 0x1B, 0x1B)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size = Pt(17)
    r.font.bold = True
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Equipo SafeAlert  ·  06 de septiembre de 2026  ·  "
                  "Documentacion generada a partir del codigo fuente")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


def anadir_control_documento(doc, filas_extra=None):
    doc.add_heading("Control de documento", level=1)
    filas = [
        ["Proyecto", "SafeAlert"],
        ["Documento", "Documentacion Tecnica Integral"],
        ["Version", "1.0"],
        ["Fecha", "06/09/2026"],
        ["Estado", "Generado"],
        ["Fuente", "Analisis del codigo fuente del repositorio"],
        ["Confidencialidad", "No incluye secretos; valores sensibles ocultos"],
    ]
    if filas_extra:
        filas.extend(filas_extra)
    md_docx._anadir_tabla(doc, filas)


def anadir_tabla_contenidos(doc):
    doc.add_heading("Indice de contenido", level=1)
    p = doc.add_paragraph()
    r = p.add_run("(Para actualizar el indice en Word: seleccionar el indice y "
                  "pulsar F9, o Herramientas > Actualizar campos).")
    r.font.size = Pt(9)
    r.font.italic = True
    parrafo = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Indice"
    run.append(t)
    fld.append(run)
    parrafo._p.append(fld)
    doc.add_page_break()


def guardar_y_verificar(doc, ruta):
    doc.save(ruta)
    d2 = Document(ruta)
    print("OK %s (parrafos=%d tablas=%d)" %
          (os.path.basename(ruta), len(d2.paragraphs), len(d2.tables)))
    return len(d2.paragraphs)


# ---------------------------------------------------------------------------
# Construccion de tablas a partir de JSON
# ---------------------------------------------------------------------------

def tabla_funciones():
    data = leer_json("funciones.json") or []
    if not data:
        return "_(sin datos)_"
    filas = sorted(data, key=lambda x: (x["archivo"], x["linea"]))
    cuerpo = [["Archivo", "Linea", "Nombre", "Lenguaje"]]
    for f in filas:
        cuerpo.append([f["archivo"], str(f["linea"]), f["nombre"],
                       f["lenguaje"]])
    return tabla_markdown(cuerpo[0], cuerpo[1:])


def tabla_clases():
    data = leer_json("clases.json") or []
    cuerpo = [["Archivo", "Linea", "Nombre", "Lenguaje"]]
    for f in sorted(data, key=lambda x: (x["archivo"], x["linea"])):
        cuerpo.append([f["archivo"], str(f["linea"]), f["nombre"],
                       f["lenguaje"]])
    return tabla_markdown(cuerpo[0], cuerpo[1:]) if len(cuerpo) > 1 \
        else "_(sin datos)_"


def tabla_tipos():
    data = leer_json("tipos.json") or []
    cuerpo = [["Archivo", "Linea", "Nombre", "Tipo"]]
    for f in sorted(data, key=lambda x: (x["archivo"], x["linea"])):
        cuerpo.append([f["archivo"], str(f["linea"]), f["nombre"], f["tipo"]])
    return tabla_markdown(cuerpo[0], cuerpo[1:]) if len(cuerpo) > 1 \
        else "_(sin datos)_"


def tabla_endpoints():
    data = leer_json("endpoints.json") or []
    cuerpo = [["Metodo/Trigger", "Ruta / Recurso", "Funcion", "Archivo",
               "Linea"]]
    for f in sorted(data, key=lambda x: (x["archivo"], x["linea"])):
        cuerpo.append([f["metodo"], f["ruta"], f["funcion"], f["archivo"],
                       str(f["linea"])])
    return tabla_markdown(cuerpo[0], cuerpo[1:]) if len(cuerpo) > 1 \
        else "_(sin datos)_"


def tabla_dependencias():
    data = leer_json("dependencias.json") or []
    cuerpo = [["Origen / Manifest", "Paquete", "Version", "Tipo"]]
    for f in sorted(data, key=lambda x: (x["origen"], x["nombre"].lower())):
        cuerpo.append([f["origen"], f["nombre"], f["version"], f["tipo"]])
    return tabla_markdown(cuerpo[0], cuerpo[1:]) if len(cuerpo) > 1 \
        else "_(sin datos)_"


def tabla_entorno():
    data = leer_json("entorno.json") or []
    cuerpo = [["Variable", "Archivo", "Contiene valor real en el repo"]]
    for f in sorted(data, key=lambda x: (x["archivo"], x["variable"])):
        cuerpo.append([f["variable"], f["archivo"],
                       "SI (oculto)" if f["valor_real_presente"] else "NO"])
    return tabla_markdown(cuerpo[0], cuerpo[1:]) if len(cuerpo) > 1 \
        else "_(sin datos)_"


def seccion_bd():
    data = leer_json("database.json") or []
    if not data:
        return "_(sin datos)_"
    partes = []
    for t in data:
        partes.append("### Tabla `%s` (archivo: %s)" % (t["tabla"],
                                                         t["archivo"]))
        enc = ["Campo", "Tipo", "Null", "Clave", "Default", "Descripcion"]
        filas = []
        for c in t["columnas"]:
            claves = []
            if c.get("pk"):
                claves.append("PK")
            if c.get("unique"):
                claves.append("UQ")
            if c.get("auto"):
                claves.append("AI")
            filas.append([c["campo"], c["tipo"],
                          c["null"], " ".join(claves) or "-",
                          c.get("default") or "-", c.get("descripcion") or ""])
        partes.append(tabla_markdown(enc, filas))
    return "\n\n".join(partes)


def tabla_resumen_directorios():
    ruta_inv = os.path.join(INT, "inventario", "inventario.json")
    if not os.path.exists(ruta_inv):
        return "_(sin datos)_"
    with open(ruta_inv, "r", encoding="utf-8") as f:
        archivos = json.load(f)
    from collections import Counter
    por_dir = {}
    for a in archivos:
        d = a["directorio"] or "(raiz)"
        por_dir.setdefault(d, Counter())["total"] += 1
        por_dir[d][a["categoria"]] += 1
    filas = []
    for d in sorted(por_dir):
        c = por_dir[d]
        desc = ", ".join("%d %s" % (c[k], k.lower())
                         for k in c if k != "total")
        filas.append([d, str(c["total"]), desc])
    return tabla_markdown(["Directorio", "Archivos", "Categorias"], filas)


def _leer_mds_todos():
    """Devuelve lista de (archivo_md, texto) bajo intermediate/modulos."""
    resultado = []
    patron = os.path.join(INT, "modulos", "*", "*.md")
    for ruta in sorted(glob.glob(patron)):
        with open(ruta, "r", encoding="utf-8") as f:
            resultado.append((os.path.basename(ruta), f.read()))
    return resultado


def _primer_titulo_archivo(texto):
    m = re.search(r"^# Archivo: (.+)$", texto, re.MULTILINE)
    if m:
        return m.group(1).strip()
    m = re.search(r"^# Documento: (.+)$", texto, re.MULTILINE)
    return m.group(1).strip() if m else "(documento)"


def _extraer_hallazgos(mds):
    """Devuelve filas (archivo, severidad, texto) de hallazgos de seguridad."""
    filas = []
    patron = re.compile(
        r"^\s*[-*]?\s*\[(CRÍTICO|CRITICO|ALTO|MEDIO|BAJO|INFORMATIVO)\]\s*"
        r":?\s*(.+)$")
    for nombre, texto in mds:
        archivo = _primer_titulo_archivo(texto)
        en_seguridad = False
        for linea in texto.split("\n"):
            if linea.startswith("## Seguridad"):
                en_seguridad = True
                continue
            if linea.startswith("## "):
                en_seguridad = False
                continue
            if not en_seguridad:
                continue
            m = patron.match(linea)
            if m:
                severidad = m.group(1).upper().replace("CRITICO", "CRÍTICO")
                filas.append([archivo, severidad, m.group(2).strip()])
    return filas


def _extraer_observaciones(mds):
    filas = []
    patron = re.compile(
        r"^\s*[-*]?\s*\[(OBSERVACIÓN TÉCNICA|POTENCIALMENTE NO UTILIZADO|"
        r"CÓDIGO LEGADO|PARCIALMENTE IMPLEMENTADA)\]\s*:?\s*(.+)$")
    for nombre, texto in mds:
        archivo = _primer_titulo_archivo(texto)
        for linea in texto.split("\n"):
            m = patron.match(linea)
            if m:
                filas.append([archivo, m.group(1), m.group(2).strip()])
    return filas


def tabla_seguridad():
    filas = _extraer_hallazgos(_leer_mds_todos())
    if not filas:
        return "_(sin hallazgos extraidos)_"
    orden = {"CRÍTICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3, "INFORMATIVO": 4}
    filas.sort(key=lambda f: (orden.get(f[1], 9), f[0].lower()))
    return tabla_markdown(["Archivo", "Severidad", "Hallazgo"], filas,
                          max_celda=260)


def tabla_observaciones():
    filas = _extraer_observaciones(_leer_mds_todos())
    if not filas:
        return "_(sin observaciones extraidas)_"
    filas.sort(key=lambda f: (f[0].lower(), f[1]))
    return tabla_markdown(["Archivo", "Marcador", "Observación"], filas,
                          max_celda=260)


def token_arbol():
    ruta = os.path.join(INT, "inventario", "arbol.txt")
    if not os.path.exists(ruta):
        return "_(sin inventario)_"
    with open(ruta, "r", encoding="utf-8") as f:
        return "```text\n" + f.read() + "```"


def token_estadisticas():
    est = leer_json("estadisticas.json")
    if not est:
        return "_(sin estadisticas)_"
    excl = est.get("directorios_excluidos_conteo", {})
    linea_excl = "; ".join("%s: %d" % (k, v) for k, v in excl.items())
    # Numero de md de analisis generados
    n_md = len(glob.glob(os.path.join(INT, "modulos", "*", "*.md")))
    n_inv = est.get("total_archivos_inventariados", 0)
    partes = [
        "| Metrica | Valor |",
        "| --- | --- |",
        "| Archivos inventariados (relevantes) | %d |" % n_inv,
        "| Archivos de analisis .md generados | %d |" % n_md,
        "| Directorios excluidos (solo conteo) | %s |" % linea_excl,
    ]
    return "\n".join(partes)


def reemplazar_tokens(texto):
    generadores = {
        "[[FUNCIONES_TABLA]]": tabla_funciones,
        "[[CLASES_TABLA]]": tabla_clases,
        "[[TIPOS_TABLA]]": tabla_tipos,
        "[[ENDPOINTS_TABLA]]": tabla_endpoints,
        "[[DEPENDENCIAS_TABLA]]": tabla_dependencias,
        "[[ENV_TABLA]]": tabla_entorno,
        "[[BD_TABLA]]": seccion_bd,
        "[[RESUMEN_DIRECTORIOS_TABLA]]": tabla_resumen_directorios,
        "[[HALLAZGOS_SEGURIDAD_TABLA]]": tabla_seguridad,
        "[[OBSERVACIONES_TABLA]]": tabla_observaciones,
        "[[ARBOL_INVENTARIO]]": token_arbol,
        "[[ESTADISTICAS_FINALES]]": token_estadisticas,
    }
    for token, fn in generadores.items():
        if token in texto:
            texto = texto.replace(token, fn())
    return texto


def construir_anexos():
    os.makedirs(ANEXOS_DIR, exist_ok=True)
    total_md = 0
    for anexo in ANEXOS:
        doc = Document()
        md_docx.cabecera_pie(doc, "SafeAlert — " + anexo["codigo"])
        anadir_portada_anexo(doc, anexo["titulo"],
                             anexo["carpetas"])
        md_docx._configurar_estilos(doc)
        md_docx.anadir_markdown(doc, anexo["intro"])
        mds = []
        for carpeta in anexo["carpetas"]:
            patron = os.path.join(INT, "modulos", carpeta, "*.md")
            mds.extend(sorted(glob.glob(patron)))
        if not mds:
            p = doc.add_paragraph()
            r = p.add_run("[ADVERTENCIA] No se encontraron archivos de "
                          "analisis para este anexo.")
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        for ruta_md in mds:
            try:
                with open(ruta_md, "r", encoding="utf-8") as f:
                    texto = f.read()
                md_docx.anadir_markdown(doc, texto)
                total_md += 1
            except Exception as exc:  # noqa: BLE001 - tolerancia por archivo
                print("ERROR convirtiendo %s: %r" % (ruta_md, exc))
                p = doc.add_paragraph()
                r = p.add_run("[ADVERTENCIA] Error de conversion del analisis "
                              "%s: %s" % (os.path.basename(ruta_md), exc))
                r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        ruta_salida = os.path.join(ANEXOS_DIR, anexo["archivo"])
        guardar_y_verificar(doc, ruta_salida)
    print("Anexos generados; archivos md procesados:", total_md)


def construir_maestro():
    doc = Document()
    md_docx.cabecera_pie(doc, "SafeAlert")
    anadir_portada_maestro(
        doc,
        "DOCUMENTACION TECNICA COMPLETA",
        "Documentacion tecnica, funcional y operativa del codigo fuente")
    md_docx._configurar_estilos(doc)
    anadir_control_documento(doc, [
        ["Alcance", "Documento maestro + anexos A-G"],
        ["Cobertura", "Archivos relevantes del repositorio"],
    ])
    anadir_tabla_contenidos(doc)

    mds = sorted(glob.glob(os.path.join(MASTER_DIR, "*.md")))
    if not mds:
        print("ADVERTENCIA: sin contenido maestro en", MASTER_DIR)
    for ruta_md in mds:
        try:
            with open(ruta_md, "r", encoding="utf-8") as f:
                texto = f.read()
            texto = reemplazar_tokens(texto)
            md_docx.anadir_markdown(doc, texto)
        except Exception as exc:  # noqa: BLE001 - tolerancia por archivo
            print("ERROR convirtiendo maestro %s: %r" % (ruta_md, exc))
            p = doc.add_paragraph()
            r = p.add_run("[ADVERTENCIA] Error de conversion del capitulo "
                          "%s: %s" % (os.path.basename(ruta_md), exc))
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    ruta_salida = os.path.join(OUT,
                               "DOCUMENTACION_TECNICA_COMPLETA_SAFEALERT.docx")
    guardar_y_verificar(doc, ruta_salida)


def main():
    global ANEXOS
    args = sys.argv[1:]
    solo = None
    letras = None
    if args:
        solo = args[0]
        if len(args) > 1:
            letras = [a.upper() for a in args[1:]]
    if solo in (None, "anexos"):
        if letras:
            seleccion = []
            for anexo in ANEXOS:
                num = anexo["codigo"].split()[-1].upper()
                if num in letras:
                    seleccion.append(anexo)
            ANEXOS = seleccion
        construir_anexos()
    if solo in (None, "maestro"):
        construir_maestro()
    return 0


if __name__ == "__main__":
    sys.exit(main())
