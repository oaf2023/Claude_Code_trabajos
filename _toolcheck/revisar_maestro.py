# ============================================================================
# Archivo         : revisar_maestro.py
# Descripcion     : Comprueba tokens sin reemplazar y lista los titulos
#                   (Heading 1/2) del documento maestro DOCX.
# Autor           : Equipo SafeAlert
# Fecha           : 2026-09-06
# Version         : 1.0.0
# Lenguaje        : Python 3.13
# Uso             : python revisar_maestro.py <ruta_docx>
# ============================================================================
import re
import sys

from docx import Document

ruta = sys.argv[1]
doc = Document(ruta)
txt = []
for p in doc.paragraphs:
    txt.append(p.text)
for t in doc.tables:
    for fila in t.rows:
        for celda in fila.cells:
            txt.append(celda.text)
completo = "\n".join(txt)
tokens = re.findall(r"\[\[[A-Z_0-9]+\]\]", completo)
print("Tokens sin reemplazar:", tokens if tokens else "ninguno")

for p in doc.paragraphs:
    if p.style.name in ("Heading 1", "Heading 2"):
        print("[%s] %s" % (p.style.name, p.text[:90]))
