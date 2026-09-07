# ============================================================================
# Archivo         : prueba_docx.py
# Descripcion     : Prueba de humo del conversor md_docx con un .md real.
# Autor           : Equipo SafeAlert
# Fecha           : 2026-09-06
# Version         : 1.0.0
# Lenguaje        : Python 3.13
# Uso             : python prueba_docx.py <md_entrada> <docx_salida>
# ============================================================================
import sys
from docx import Document

sys.path.insert(0, "C:\\Claude_Code_trabajos\\_toolcheck")
import md_docx  # noqa: E402

entrada = sys.argv[1]
salida = sys.argv[2]

doc = Document()
md_docx.cabecera_pie(doc, "SafeAlert — Prueba")
with open(entrada, "r", encoding="utf-8") as f:
    texto = f.read()
md_docx.anadir_markdown(doc, texto)
doc.save(salida)

# Reabrir y contar
doc2 = Document(salida)
parrafos = len(doc2.paragraphs)
tablas = len(doc2.tables)
print("OK parrafos=%d tablas=%d archivo=%s" % (parrafos, tablas, salida))
